from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from document.schema import PassageResult
from utils.text_processing import split_sentences


class DocxBuilder:
    """DOCX 문서 생성기. 기존 HWP 템플릿 레이아웃을 DOCX로 재현."""

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        """A4 페이지, 여백 설정"""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    def _setup_styles(self):
        """기본 스타일 설정"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(10)
        # 한글 폰트 설정
        rpr = style.element.get_or_add_rPr()
        ea_font = rpr.makeelement(qn('w:rFonts'), {})
        ea_font.set(qn('w:eastAsia'), '맑은 고딕')
        rpr.append(ea_font)

    def _add_heading(self, text: str, level: int = 1):
        """제목 추가"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(50, 50, 50)
        p.space_after = Pt(6)
        return p

    def _add_section_title(self, title: str):
        """섹션 구분 제목 (FLOW, VOCA, T/F, SUMMARY)"""
        p = self.doc.add_paragraph()
        p.space_before = Pt(12)
        run = p.add_run(f"■ {title}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 30, 120)
        p.space_after = Pt(4)

    def _set_cell_font(self, cell, text: str, size: int = 9,
                       bold: bool = False, color: RGBColor = None):
        """테이블 셀 텍스트 및 서식 설정"""
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        # 셀 내부 여백
        p.space_before = Pt(1)
        p.space_after = Pt(1)

    def _set_table_borders(self, table):
        """테이블 테두리 설정"""
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
        borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            element = borders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '0',
                qn('w:color'): '999999',
            })
            borders.append(element)
        tbl_pr.append(borders)

    def add_passage(self, data: PassageResult):
        """지문 1개의 전체 레이아웃 추가"""
        d = data.to_dict()

        # 1. 헤더: 지문 번호 + 주제
        self._add_heading(f"[{d.get('No', '')}] {d.get('Topic', '')}")

        # 2. 본문 테이블 (번호 / 영어 / 한국어)
        self._add_eng_kor_table(d.get('Eng', ''), d.get('Kor', ''))

        # 3. FLOW 섹션
        self._add_flow_section(d.get('Flow', ''))

        # 4. VOCA 섹션
        self._add_voca_section(data)

        # 5. T/F 섹션
        self._add_tf_section(d.get('TF', ''), d.get('TFA', ''))

        # 6. SUMMARY 섹션
        self._add_summary_section(d.get('TST', ''))

        # 페이지 구분
        self.doc.add_page_break()

    def _add_eng_kor_table(self, eng_text: str, kor_text: str):
        """영어-한국어 대조 테이블"""
        self._add_section_title("본문 (TEXT)")

        eng_sentences = split_sentences(eng_text)
        kor_sentences = split_sentences(kor_text)
        row_count = max(len(eng_sentences), len(kor_sentences))

        if row_count == 0:
            return

        table = self.doc.add_table(rows=row_count, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)

        # 열 너비 설정 (번호:영어:한국어 = 0.8:7.5:7.5)
        for row in table.rows:
            row.cells[0].width = Cm(0.8)
            row.cells[1].width = Cm(7.5)
            row.cells[2].width = Cm(7.5)

        for i in range(row_count):
            # 번호
            self._set_cell_font(table.cell(i, 0), str(i + 1),
                                size=8, bold=True,
                                color=RGBColor(100, 100, 100))
            table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 영어
            eng = eng_sentences[i] if i < len(eng_sentences) else ""
            self._set_cell_font(table.cell(i, 1), eng, size=9)

            # 한국어
            kor = kor_sentences[i] if i < len(kor_sentences) else ""
            self._set_cell_font(table.cell(i, 2), kor, size=9,
                                color=RGBColor(80, 80, 80))

    def _add_flow_section(self, flow_text: str):
        """흐름 분석 섹션"""
        self._add_section_title("흐름 분석 (FLOW)")
        if not flow_text:
            return
        for line in flow_text.split('\n'):
            line = line.strip()
            if line:
                p = self.doc.add_paragraph()
                run = p.add_run(f"  {line}")
                run.font.size = Pt(9)
                p.space_after = Pt(2)

    def _add_voca_section(self, data: PassageResult):
        """핵심 어휘 섹션 (3열 테이블)"""
        self._add_section_title("핵심 어휘 (VOCA)")

        voca_list = data.get_voca_list()
        if not voca_list:
            return

        # 3열로 배치 (6행 × 3열 = 최대 18개)
        cols = 3
        rows = (len(voca_list) + cols - 1) // cols
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)

        for idx, voca in enumerate(voca_list):
            row_idx = idx // cols
            col_idx = idx % cols
            self._set_cell_font(table.cell(row_idx, col_idx), voca, size=9)

    def _add_tf_section(self, tf_text: str, tfa_text: str):
        """T/F 문제 섹션"""
        self._add_section_title("True / False")

        if not tf_text:
            return

        for line in tf_text.split('\n'):
            line = line.strip()
            if line:
                p = self.doc.add_paragraph()
                run = p.add_run(f"  {line}")
                run.font.size = Pt(9)
                p.space_after = Pt(2)

        # 정답
        if tfa_text:
            p = self.doc.add_paragraph()
            p.space_before = Pt(6)
            run = p.add_run(f"  정답: {tfa_text}")
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor(180, 0, 0)

    def _add_summary_section(self, tst_text: str):
        """영어 요약 + 한국어 해석 섹션"""
        self._add_section_title("요약 (SUMMARY)")

        if not tst_text:
            return

        lines = tst_text.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if line:
                p = self.doc.add_paragraph()
                run = p.add_run(f"  {line}")
                run.font.size = Pt(9)
                # 한국어 해석(두 번째 줄 이후)은 회색으로 표시
                if i > 0:
                    run.font.color.rgb = RGBColor(80, 80, 80)
                p.space_after = Pt(2)

    def build(self) -> bytes:
        """DOCX 파일을 바이트로 반환 (st.download_button 호환)"""
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
